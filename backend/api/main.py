import os
import io
import logging
import traceback
from typing import List
from datetime import datetime

import requests
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from sqlalchemy import create_engine, Column, Integer, Text, DateTime, text
from sqlalchemy.orm import sessionmaker, declarative_base
from sqlalchemy.dialects.postgresql import JSONB
from pgvector.sqlalchemy import Vector
from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import docx
from PIL import Image

# Vercel-specific import
from mangum import Mangum

# ----------------------------
# Load environment variables
# ----------------------------
load_dotenv()

DB_USER = os.getenv("DB_USER", "dev_admin")
DB_PASSWORD = os.getenv("DB_PASSWORD", "password")
DB_HOST = os.getenv("DB_HOST", "localhost")
DB_PORT = os.getenv("DB_PORT", "5432")
DB_NAME = os.getenv("DB_NAME", "sop_dev_db")

DATABASE_URL = (
    f"postgresql+psycopg2://{DB_USER}:{DB_PASSWORD}@{DB_HOST}:{DB_PORT}/{DB_NAME}?sslmode=require"
)

OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3:latest")

# ----------------------------
# Database setup
# ----------------------------
engine = create_engine(DATABASE_URL, echo=True, pool_pre_ping=True)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

class Document(Base):
    __tablename__ = "documents"
    id = Column(Integer, primary_key=True, index=True)
    content = Column(Text, nullable=False)
    doc_metadata = Column(JSONB, nullable=True)
    embedding = Column(Vector(768))

class ChatHistory(Base):
    __tablename__ = "chat_history"
    id = Column(Integer, primary_key=True, index=True)
    user_query = Column(Text, nullable=False)
    answer = Column(Text, nullable=False)
    timestamp = Column(DateTime, default=datetime.utcnow)

Base.metadata.create_all(bind=engine)

# ----------------------------
# Embedding model
# ----------------------------
EMBED_MODEL = SentenceTransformer("all-mpnet-base-v2")

def embed_texts(texts: List[str]) -> List[List[float]]:
    embeddings = EMBED_MODEL.encode(texts, convert_to_numpy=True, show_progress_bar=False)
    return [emb.tolist() for emb in embeddings]

# ----------------------------
# FastAPI setup
# ----------------------------
app = FastAPI(title="SOP Chatbot (Ollama + PostgreSQL + pgvector)")

# logger for diagnostics
logger = logging.getLogger("sop_upload")
logger.setLevel(logging.INFO)

# Enable CORS for frontend communication
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # For dev; restrict in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ----------------------------
# Request/Response models
# ----------------------------
class ChatRequest(BaseModel):
    query: str

# ----------------------------
# Root route & Health check
# ----------------------------
@app.get("/")
def root():
    return {"message": "SOP Chatbot backend is running. Use /api/chat or /api/upload."}

@app.get("/health")
def health_check():
    db = SessionLocal()
    try:
        db.execute(text("SELECT 1"))
        return {"status": "ok", "db": "connected"}
    except Exception as e:
        return {"status": "error", "db": "disconnected", "error": str(e)}
    finally:
        db.close()

# ----------------------------
# File Upload Endpoint
# ----------------------------
@app.post("/api/upload")
async def upload(file: UploadFile = File(...)):
    filename = file.filename
    ext = os.path.splitext(filename)[1].lower()
    if not ext:
        ct = getattr(file, "content_type", "") or ""
        ct = ct.lower()
        if ct.startswith("image/"):
            ext = "." + ct.split("/")[-1]
        elif "pdf" in ct:
            ext = ".pdf"
        elif "wordprocessingml" in ct or "officedocument.wordprocessingml" in ct:
            ext = ".docx"
        elif "msword" in ct:
            ext = ".doc"
    extracted_text = ""
    file_bytes = await file.read()
    logger.info(f"Upload received: filename={filename!r} ext={ext!r} content_type={getattr(file, 'content_type', None)!r} size={len(file_bytes)} bytes")
    try:
        if ext in [".txt", ".md"]:
            extracted_text = file_bytes.decode("utf-8", errors="ignore")
        elif ext == ".pdf":
            pdf = PdfReader(io.BytesIO(file_bytes))
            for page in pdf.pages:
                extracted_text += page.extract_text() or ""
        elif ext == ".docx":
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
        elif ext == ".doc":
            try:
                import textract
                import tempfile
                with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tf:
                    tf.write(file_bytes)
                    temp_path = tf.name
                try:
                    extracted = textract.process(temp_path, extension="doc")
                    extracted_text = extracted.decode("utf-8", errors="ignore")
                finally:
                    try:
                        os.unlink(temp_path)
                    except Exception:
                        pass
            except Exception as tex_err:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "Unsupported .doc processing: textract not available or failed. "
                        "Please convert the .doc file to .docx or PDF and try again, or install textract with its dependencies. "
                        f"(textract error: {str(tex_err)})"
                    ),
                )
        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
            image = Image.open(io.BytesIO(file_bytes))
            extracted_text = f"Image file: {filename}, format: {image.format}, size: {image.size}"
        else:
            supported = ".txt, .md, .pdf, .docx, .png, .jpg, .jpeg, .bmp, .gif, .webp"
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type '{ext}' for file '{filename}'. Supported types: {supported}",
            )
    except Exception as e:
        tb = traceback.format_exc()
        logger.error(f"Error processing file {filename}: {str(e)}\n{tb}")
        raise HTTPException(status_code=400, detail=f"File processing error: {str(e)}")

    if not extracted_text or not extracted_text.strip():
        if ext == ".pdf":
            detail = (
                "No text content found in PDF. This may be a scanned/image PDF or a PDF with non-extractable text. "
                "Install optional tools to improve extraction: `pip install pdfplumber` for better PDF parsing, or "
                "install Tesseract OCR and `pip install pytesseract` to OCR image PDFs. Alternatively convert the PDF to a text-based PDF or to DOCX and re-upload."
            )
        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
            detail = (
                "No text found in image. If this is a scanned image containing text, install Tesseract OCR (system package) and `pip install pytesseract` to enable OCR."
            )
        else:
            detail = "No text content found in file."
        raise HTTPException(status_code=400, detail=detail)

    if ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
        chunks = [extracted_text]
    else:
        chunks = [p.strip() for p in extracted_text.split("\n\n") if p.strip()]

    if not chunks:
        logger.info(f"No text content extracted for {filename}. ext={ext} extracted_text_len={len(extracted_text)}")
        if ext == ".pdf":
            detail = (
                "No text content found in PDF. This may be a scanned/image PDF or a PDF with non-extractable text. "
                "Install optional tools to improve extraction: `pip install pdfplumber` for better PDF parsing, or "
                "install Tesseract OCR and `pip install pytesseract` to OCR image PDFs. Alternatively convert the PDF to a text-based PDF or to DOCX and re-upload."
            )
        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
            detail = (
                "No text found in image. If this is a scanned image containing text, install Tesseract OCR (system package) and `pip install pytesseract` to enable OCR."
            )
        else:
            detail = "No text content found in file."
        raise HTTPException(status_code=400, detail=detail)

    embeds = embed_texts(chunks)

    db = SessionLocal()
    try:
        for c, e in zip(chunks, embeds):
            doc = Document(content=c, doc_metadata={"source": filename}, embedding=e)
            db.add(doc)
        db.commit()
    finally:
        db.close()

    return {"status": "ok", "filetype": ext, "chunks": len(chunks)}

# ----------------------------
# Chat Endpoint
# ----------------------------
@app.post("/api/chat")
async def chat(request: ChatRequest):
    query = request.query
    query_emb = embed_texts([query])[0]

    db = SessionLocal()
    try:
        emb_str = ",".join(str(x) for x in query_emb)
        sql = text(
            f"""
            SELECT id, content, doc_metadata, embedding <#> ARRAY[{emb_str}]::vector AS distance
            FROM documents
            ORDER BY embedding <#> ARRAY[{emb_str}]::vector ASC
            LIMIT 3
            """
        )
        results = db.execute(sql).fetchall()
        context_docs = [row[1] for row in results]

        context = "\n---\n".join(context_docs)
        prompt = f"""You are a helpful assistant. Use the following context to answer the user's question.

Context:
{context}

Question: {query}
Answer:"""

        answer = None
        last_exception = None
        max_attempts = 3
        timeout_seconds = 600
        for attempt in range(1, max_attempts + 1):
            try:
                resp = requests.post(
                    f"{OLLAMA_URL}/api/generate",
                    json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": False},
                    timeout=timeout_seconds,
                )
                resp.raise_for_status()
                try:
                    data = resp.json()
                except Exception:
                    raw = resp.text
                    answer = f"LLM returned non-JSON response: {raw[:1000]}"
                    break

                if isinstance(data, dict):
                    if "response" in data and data["response"]:
                        answer = data["response"]
                    elif "answer" in data and data["answer"]:
                        answer = data["answer"]
                    elif "choices" in data and isinstance(data["choices"], list):
                        parts = []
                        for c in data["choices"]:
                            if isinstance(c, dict):
                                if "text" in c and c["text"]:
                                    parts.append(c["text"])
                                elif "message" in c and isinstance(c["message"], dict):
                                    cont = c["message"].get("content") or c["message"].get("text")
                                    if cont:
                                        parts.append(cont)
                        answer = "\n".join(parts).strip() if parts else None
                    elif "text" in data and data["text"]:
                        answer = data["text"]
                if not answer:
                    possible = []
                    for k in ("output", "result", "completion"):
                        if k in data:
                            possible.append(str(data[k]))
                    if possible:
                        answer = " ".join(possible)
                    else:
                        answer = "No answer (unexpected response shape)."

                break
            except requests.exceptions.Timeout as e:
                last_exception = e
                if attempt < max_attempts:
                    import time
                    wait = attempt * 2
                    time.sleep(wait)
                    continue
                else:
                    answer = "LLM error: Request timed out."
            except requests.exceptions.ConnectionError as e:
                last_exception = e
                answer = "LLM error: Could not connect to the local model."
                break
            except Exception as e:
                last_exception = e
                answer = f"LLM error: {str(e)}"
                break

        chat_entry = ChatHistory(user_query=query, answer=answer)
        db.add(chat_entry)
        db.commit()

        return {"answer": answer, "context": context_docs}
    finally:
        db.close()

# Vercel handler for serverless deployment
handler = Mangum(app)
